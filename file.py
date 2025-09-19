from openpyxl import Workbook

from openpyxl.styles import Font, Alignment

from docx import Document
 
# Create Excel file

wb = Workbook()

ws = wb.active

ws.title = "Panel Evaluation Grid"
 
# Define headers

headers = [

    "Applicant Name", 

    "1. Engagement with Prior Sessions (Q2 & Q3)", 

    "2. Cultural Learning Interest & Co-design Belief (Q4)", 

    "3. Project Relevance & Co-design Readiness (Q5 & Q6)", 

    "4. Workshop Fit: Open-mindedness, Creativity, Presence (Inferred)", 

    "Total Score (out of 20)", 

    "Panel Notes", 

    "Unique Perspective? (Y/N)"

]

ws.append(headers)
 
# Style header row

for cell in ws[1]:

    cell.font = Font(bold=True)

    cell.alignment = Alignment(horizontal='center', vertical='center')
 
# Set column widths for better readability

col_widths = [20, 15, 15, 15, 20, 15, 30, 20]

for i, width in enumerate(col_widths, 1):

    ws.column_dimensions[chr(64+i)].width = width
 
# Save Excel file

wb.save("Workshop_Participant_Evaluation_Grid.xlsx")
 
# Create Word document

doc = Document()

doc.add_heading('Workshop Participant Evaluation Grid (Panel Use)', level=1)
 
# Add table

table = doc.add_table(rows=2, cols=8)

table.style = 'Table Grid'

headers_word = [

    "Applicant Name",

    "Engagement\n(Q2 & Q3)",

    "Interest in Co-design\n(Q4)",

    "Project Relevance\n(Q5 & Q6)",

    "Workshop Fit\n(Q4, Q5)",

    "Total\n(out of 20)",

    "Panel Notes",

    "Unique Perspective?\n(Y/N)"

]

hdr_cells = table.rows[0].cells

for i, header in enumerate(headers_word):

    hdr_cells[i].text = header
 
# Leave second row empty for input

for cell in table.rows[1].cells:

    cell.text = ""
 
# Save Word document

doc.save("Workshop_Participant_Evaluation_Grid.docx")

 
